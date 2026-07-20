from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- PATH CONFIGURATION ---
BASE_DIR = r"C:\Users\SG\.gemini\antigravity\brain\55e750c7-2f22-4c7f-9088-7253e0f174af"
IMG = {
    "dashboard":     os.path.join(BASE_DIR, "demo_dashboard_1775150954601.png"),
    "leads":         os.path.join(BASE_DIR, "demo_leads_1775151003113.png"),
    "deals":         os.path.join(BASE_DIR, "demo_deals_1775151044116.png"),
    "performance":   os.path.join(BASE_DIR, "demo_performance_1775151078973.png"),
    "organizations": os.path.join(BASE_DIR, "demo_organizations_1775151115654.png"),
    "contacts":      os.path.join(BASE_DIR, "demo_contacts_1775151149991.png"),
    "automations":   os.path.join(BASE_DIR, "demo_automations_1775151210912.png"),
}

PRIMARY   = RGBColor(30, 27, 75)    # Navy Indigo
ACCENT    = RGBColor(79, 70, 229)   # Bright Indigo
GREY      = RGBColor(107, 114, 128) # Muted grey
LIGHT     = RGBColor(243, 244, 246) # Light bg
SUCCESS   = RGBColor(5, 150, 105)   # Green

# ───────────────────────────────────────────
def shade_cell(cell, hex_color="1E1B4B"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_img(doc, key, width=6.3, caption=None):
    path = IMG.get(key, "")
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(f"Figure: {caption}")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size    = Pt(9)
        cp.runs[0].italic       = True
        cp.runs[0].font.color.rgb = GREY

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = PRIMARY
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = ACCENT
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = GREY
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Pt(20 * (level + 1))
    return p

def numbered(doc, text):
    return doc.add_paragraph(text, style='List Number')

def tip_box(doc, text, label="💡 Pro Tip"):
    """Light callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(f"{label}  {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = SUCCESS
    return p

def toc_row(doc, section, title, pg):
    p  = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{section}  {title}")
    r1.font.size = Pt(11)
    # Right-align page number using tab stop
    r2 = p.add_run(f"\t{pg}")
    r2.font.size = Pt(11)
    r2.bold = True
    r2.font.color.rgb = ACCENT

def divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(220, 220, 220)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ═══════════════════════════════════════════
def create_manual():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Global font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── COVER PAGE ──────────────────────────────────────────
    for _ in range(4): doc.add_paragraph("")

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("DigiXCrm CRM")
    r.bold = True
    r.font.size = Pt(52)
    r.font.color.rgb = PRIMARY

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run("Professional User & Administration Guide")
    r.italic = True
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run("Enterprise Release v2.4  ·  DigiXCrm Global HQ Demo Environment")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    doc.add_paragraph("")
    add_img(doc, "dashboard", 6.5, "DigiXCrm Executive Dashboard — Live Demo Environment")
    doc.add_page_break()

    # ── TABLE OF CONTENTS ────────────────────────────────────
    h1(doc, "Table of Contents")
    divider(doc)
    toc_entries = [
        ("1.",  "Executive Introduction & Platform Overview",          "3"),
        ("2.",  "System Architecture & Security Model",                "4"),
        ("3.",  "Getting Started: Onboarding & First Login",           "5"),
        ("4.",  "Workspace Hub & Multi-Tenant Navigation",             "6"),
        ("5.",  "Executive Dashboard — Real-Time Intelligence",        "7"),
        ("6.",  "Lead Management & Acquisition",                       "8"),
        ("  6.1", "Creating a Lead Manually",                          "8"),
        ("  6.2", "Bulk CSV Import",                                   "8"),
        ("  6.3", "Lead Filtering, Search & Sorting",                  "9"),
        ("  6.4", "Lead Assignment & Notifications",                   "9"),
        ("7.",  "Sales Pipeline & Visual Deal Board",                  "10"),
        ("  7.1", "Kanban Board Overview",                             "10"),
        ("  7.2", "Creating & Managing Deals",                         "10"),
        ("  7.3", "Forecast View",                                     "11"),
        ("8.",  "Orders & Fulfillment Management",                     "11"),
        ("9.",  "Contacts Management",                                 "12"),
        ("10.", "Organizations & Account Management",                  "12"),
        ("11.", "Performance Hub & Quota Tracking",                    "13"),
        ("12.", "Activities & Interaction History",                    "14"),
        ("13.", "Global Search",                                       "14"),
        ("14.", "Workflow Automations",                                "15"),
        ("15.", "System Settings & Enterprise Configuration",          "16"),
        ("  15.1", "My Profile",                                       "16"),
        ("  15.2", "Team & Permissions (RBAC)",                        "16"),
        ("  15.3", "Integrations Panel",                               "17"),
        ("  15.4", "Data Export",                                      "17"),
        ("16.", "SaaS Administration — Control Room",                  "18"),
        ("17.", "Troubleshooting & FAQ",                               "19"),
        ("18.", "Glossary of Terms",                                   "20"),
    ]
    for sec, title, pg in toc_entries:
        toc_row(doc, sec, title, pg)
    divider(doc)
    doc.add_page_break()

    # ── SECTION 1: INTRODUCTION ───────────────────────────────
    h1(doc, "1. Executive Introduction & Platform Overview")
    body(doc,
        "DigiXCrm CRM is a high-performance, multi-tenant Sales Orchestration Platform engineered for modern "
        "enterprises that demand speed, security, and scale. The platform unifies lead capture, deal management, "
        "contact intelligence, and revenue analytics under a single, intuitive interface.")
    body(doc,
        "Whether you are a Sales Representative capturing a fresh prospect, a Manager analysing your team's "
        "pipeline health, or an Administrator configuring enterprise-grade automations — DigiXCrm provides "
        "purpose-built tools for every role in your organisation.")

    h2(doc, "1.1 Core Value Propositions")
    bullet(doc, "Real-Time Collaboration: Server-Sent Events (SSE) push updates to all users instantly without page refreshes.")
    bullet(doc, "Multi-Workspace Architecture: Manage multiple business units, brands, or client accounts from one login.")
    bullet(doc, "Granular Security: Role-based permissions down to individual action level (delete, export, view metrics).")
    bullet(doc, "Intelligent Automation: 'If This, Then That' workflow engine to automate repetitive tasks.")
    bullet(doc, "AI-Powered Insights: One-click AI analysis summarises your pipeline health and recommends next actions.")

    h2(doc, "1.2 Who Is This Guide For?")
    body(doc,
        "This manual serves three primary audiences:")
    bullet(doc, "Sales Representatives (Reps): Day-to-day usage of Leads, Contacts, Activities, and Deals.")
    bullet(doc, "Managers: Team oversight, quota tracking, pipeline reports, and lead assignment.")
    bullet(doc, "Administrators & Account Owners: System configuration, user provisioning, integrations, and billing.")

    doc.add_page_break()

    # ── SECTION 2: ARCHITECTURE & SECURITY ───────────────────
    h1(doc, "2. System Architecture & Security Model")
    body(doc,
        "DigiXCrm was built with a 'Security-First' philosophy. Every API endpoint, every data query, and every "
        "UI interaction is governed by a multi-layered permission model.")

    h2(doc, "2.1 Workspace Isolation")
    body(doc,
        "Data is isolated at the database query level using 'Workspace Scoping'. Every database query is "
        "automatically filtered by the user's active workspace ID. A Rep in 'Workspace A' can never — even "
        "accidentally — see leads, deals, or contacts belonging to 'Workspace B'.")

    h2(doc, "2.2 Role Hierarchy")
    body(doc, "The platform supports a four-tier role hierarchy:")
    bullet(doc, "Super Admin: Platform-wide god-mode. Manages all organisations, billing accounts, and global settings.")
    bullet(doc, "Account Owner: Manages one organisation (multiple workspaces, billing, seats).")
    bullet(doc, "Admin: Manages one workspace (users, stages, integrations, automations).")
    bullet(doc, "Manager → Rep: Operational roles with read/write access scoped to their workspace data.")

    h2(doc, "2.3 Password & Token Security")
    body(doc,
        "All passwords are hashed with Bcrypt (cost factor 10). Invitation and password-reset tokens are "
        "one-time-use and expire — invitation links after 3 days, reset links after 1 hour. Once used, "
        "the token is permanently deleted from the database, preventing replay attacks.")

    h2(doc, "2.4 Real-Time Sync")
    body(doc,
        "DigiXCrm uses a dual-layer synchronization model. Server-Sent Events (SSE) deliver instant push "
        "notifications to connected browsers. A lightweight HTTP polling fallback (every 30 seconds) ensures "
        "data consistency even when SSE connections are temporarily interrupted.")

    doc.add_page_break()

    # ── SECTION 3: GETTING STARTED ────────────────────────────
    h1(doc, "3. Getting Started: Onboarding & First Login")
    body(doc,
        "New users are added to DigiXCrm by an Administrator. The system follows an 'Invitation-Only' model — "
        "self-registration is not possible. This ensures that only authorised personnel gain platform access.")

    h2(doc, "3.1 Accepting an Invitation")
    numbered(doc, "Check your inbox for an email from DigiXCrm titled 'You have been invited to DigiXCrm Workspace'.")
    numbered(doc, "Click the 'Complete Setup' button inside the email.")
    numbered(doc, "On the Setup page, enter your Full Name, Professional Title (optional), Phone Number (optional), and a secure password.")
    numbered(doc, "Click 'Complete Setup'. You will be redirected to the Sign-In page.")
    tip_box(doc, "Invitation links are one-time-use only. If your link has expired (after 3 days), ask your Administrator to re-send the invitation.")

    h2(doc, "3.2 Signing In")
    numbered(doc, "Navigate to your DigiXCrm URL (e.g., https://yourcompany.digicareproducts.com/digisales/).")
    numbered(doc, "Enter your registered email address and password.")
    numbered(doc, "If Two-Factor Authentication (2FA) is enabled, enter your 6-digit code from the authenticator app.")
    numbered(doc, "Click 'Sign In'. You will land on the Executive Dashboard.")

    h2(doc, "3.3 Password Recovery")
    body(doc,
        "If you forget your password, click 'Forgot your password?' on the sign-in page. Enter your email address "
        "and a secure reset link will be sent. This link is valid for 1 hour and can only be used once.")

    doc.add_page_break()

    # ── SECTION 4: WORKSPACE HUB ─────────────────────────────
    h1(doc, "4. Workspace Hub & Multi-Tenant Navigation")
    body(doc,
        "The Workspace Hub, located in the top-left corner of the sidebar, is your gateway to multi-tenant navigation. "
        "Users assigned to multiple workspaces can switch between them instantly using the dropdown. "
        "All data — leads, deals, contacts, contacts, reports — updates immediately to reflect the selected workspace.")
    h2(doc, "4.1 Switching Workspaces")
    numbered(doc, "Click the workspace name shown in the top-left of the sidebar (e.g., 'DigiXCrm Global HQ').")
    numbered(doc, "A dropdown appears listing all workspaces you are a member of.")
    numbered(doc, "Click any workspace name to switch. The entire application context changes instantly.")
    tip_box(doc, "Each workspace has its own deal stages, automation rules, team members, and integrations. Changes in one workspace never affect another.")

    doc.add_page_break()

    # ── SECTION 5: DASHBOARD ──────────────────────────────────
    h1(doc, "5. Executive Dashboard — Real-Time Intelligence")
    add_img(doc, "dashboard", 6.3, "Executive Dashboard showing live KPIs, Revenue Trajectory chart, and Sales Funnel Diagnostics")
    body(doc,
        "The Executive Dashboard is the nerve centre of DigiXCrm. It provides an 'at-a-glance' view of your "
        "organisation's commercial health, updated in real-time as your team works.")

    h2(doc, "5.1 KPI Cards")
    bullet(doc, "Total Captured Revenue: The aggregate value of all WON deals and ONBOARDED orders within the selected period.")
    bullet(doc, "Active Deals: The current number of deals actively progressing through your pipeline. Sub-text shows total pipeline value.")
    bullet(doc, "Conversion Rate: The percentage of total leads that have been promoted to a deal. Calculated as (Converted Leads ÷ Total Leads) × 100.")
    bullet(doc, "Activities Logged: Total cross-channel interactions (calls, notes, emails) recorded by the team.")

    h2(doc, "5.2 Revenue Trajectory Chart")
    body(doc,
        "This time-series chart plots 'Pipeline Volume' (what could be won) against 'Confirmed Revenue' (what has been won). "
        "Use the period selector (top right) to change the view between 30 Days, 3 Months, and 6 Months.")

    h2(doc, "5.3 Sales Funnel Diagnostics")
    body(doc,
        "The horizontal bar chart on the right visualises your lead distribution across each status stage "
        "(New Lead, Contacted, Qualified, Converted). A healthy funnel shows a decreasing bar width from top to bottom.")

    h2(doc, "5.4 AI Analysis Report")
    body(doc,
        "Click the 'AI Analysis' button (top-right, purple) to generate an instant natural-language summary of your "
        "current performance data. The AI identifies your top performer, your weakest pipeline stage, and suggests "
        "a tactical focus for the next 30 days.")

    doc.add_page_break()

    # ── SECTION 6: LEADS ──────────────────────────────────────
    h1(doc, "6. Lead Management & Acquisition")
    add_img(doc, "leads", 6.3, "Leads Table — 250+ demo leads with status badges, source tags, and owner assignments")
    body(doc,
        "In DigiXCrm, a Lead represents any potential opportunity that has not yet been formally qualified into "
        "your sales pipeline. Every lead is tracked with rich metadata: source channel, service of interest, "
        "estimated quotation, and real-time status.")

    h2(doc, "6.1 Creating a Lead Manually")
    numbered(doc, "Click the '+ New Lead' button (top-right of the Leads page).")
    numbered(doc, "Fill in Contact Details: First Name, Last Name, Email, Phone Number.")
    numbered(doc, "Set Service Interest and enter an Estimated Quotation (optional).")
    numbered(doc, "Assign an Owner from the dropdown. The creator is set as owner by default.")
    numbered(doc, "Link to an existing Organization (optional) or create a new one inline.")
    numbered(doc, "Click 'Save Lead'. The lead appears at the top of your table instantly.")
    tip_box(doc, "Unread leads appear with a subtle blue highlight in the table. Once you click on a lead to view it, the highlight disappears, indicating it has been reviewed.")

    h2(doc, "6.2 Bulk CSV Import")
    body(doc,
        "For migrating from another CRM or uploading a large prospect list:")
    numbered(doc, "Click 'Import Leads' (top-right of the Leads page).")
    numbered(doc, "Upload your .CSV file. Required columns: firstName, lastName. Recommended: email, phone, service.")
    numbered(doc, "The system maps your CSV headers automatically. Review the mapping and correct any mismatches.")
    numbered(doc, "Click 'Import'. All leads are assigned to you as the importer by default.")
    numbered(doc, "A success notification confirms the number of leads imported. Duplicates are detected and skipped.")

    h2(doc, "6.3 Lead Filtering, Search & Sorting")
    bullet(doc, "Search Bar: Real-time search by name, email, or phone number.")
    bullet(doc, "Status Filter: Filter by New, Contacted, Qualified, or Sent to Deals.")
    bullet(doc, "Source Filter: Filter by Manual, Cold Call, Referral, Website, Trade Show, etc.")
    bullet(doc, "Date Range: Narrow down leads created within a specific time window.")
    bullet(doc, "Column Sorting: Click any column header (Date, Name, Status, Assigned To) to sort ascending or descending.")

    h2(doc, "6.4 Lead Assignment & Notifications")
    body(doc,
        "Managers and Admins can reassign any lead to a different team member using the '...' actions menu. "
        "Once reassigned, the new owner receives an in-app notification (bell icon) and an email alert with a direct "
        "link to the lead — all delivered in real-time without any page reload.")

    h2(doc, "6.5 Promoting a Lead to an Opportunity")
    body(doc,
        "When a lead is qualified and ready to move into your sales pipeline:")
    numbered(doc, "Click 'Opportunity' next to the lead in the table.")
    numbered(doc, "A deal is created automatically in the Sales Pipeline at the first stage.")
    numbered(doc, "The lead's status updates to 'Sent to Deals' in the Leads table.")

    doc.add_page_break()

    # ── SECTION 7: DEALS ──────────────────────────────────────
    h1(doc, "7. Sales Pipeline & Visual Deal Board")
    add_img(doc, "deals", 6.3, "Sales Pipeline Kanban Board — 99 active deals across stages with Total Pipeline of $13M+")
    body(doc,
        "The Sales Pipeline is where qualified opportunities are tracked, managed, and closed. DigiXCrm uses "
        "a visual Kanban board to give your team an instant picture of deal flow and pipeline health.")

    h2(doc, "7.1 Kanban Board Overview")
    body(doc,
        "Each column on the Kanban board represents a stage in your sales lifecycle. The default stages are: "
        "New Lead → Contacted → Negotiation → Sales Confirmed. Each card shows the deal title, value, "
        "assigned owner avatar, and the number of days the deal has spent in its current stage.")
    tip_box(doc, "Deals lingering in the same stage for too long are a pipeline risk. Encourage your team to update deal stages daily to ensure your pipeline data is accurate.")

    h2(doc, "7.2 Creating & Managing Deals")
    numbered(doc, "Click '+ New Deal' (top-right of the Deals page).")
    numbered(doc, "Enter the Deal Title (e.g., 'Hyperion Systems - Cloud Migration'), associated Organization, and Deal Value.")
    numbered(doc, "Set the starting Stage and assign an Owner.")
    numbered(doc, "Click 'Save'. The deal appears in the appropriate Kanban column immediately.")
    body(doc, "To update a deal's stage, click the deal card to open it and change the stage from the detail panel. The Kanban board updates in real-time for all connected team members.")

    h2(doc, "7.3 Forecast View")
    body(doc,
        "Switch from 'Kanban Board' to 'Forecast' using the tab at the top of the Deals page. The Forecast "
        "view displays a table of all open deals sorted by value, allowing you to calculate your weighted "
        "revenue forecast based on probability-adjusted deal stages.")

    h2(doc, "7.4 Import Deals")
    body(doc,
        "Similar to Lead Import, you can bulk-upload deals via CSV using the 'Import Deals' button. "
        "This is useful for migrating historical deal data from legacy CRM systems.")

    doc.add_page_break()

    # ── SECTION 8: ORDERS ─────────────────────────────────────
    h1(doc, "8. Orders & Fulfillment Management")
    body(doc,
        "When a deal is marked as 'WON' (Sales Confirmed), it transitions to the Orders section. "
        "This is a dedicated fulfillment view that separates active-selling work from post-sale delivery, "
        "allowing operations teams to manage their workload without cluttering the main sales pipeline.")

    h2(doc, "8.1 The Orders Board")
    body(doc,
        "The Orders board shows all confirmed sales awaiting or undergoing delivery. Each order card displays "
        "the deal value, the responsible owner, and the number of days since the order was confirmed.")
    tip_box(doc, "Use the Orders view as your daily checklist for fulfillment. Mark orders as 'ONBOARDED' once the client has been fully set up and delivered to.")

    h2(doc, "8.2 Closed-Lost Tracking")
    body(doc,
        "Deals that were not won are tracked in the 'Closed-Lost' section. This data is critical for coaching "
        "conversations and identifying patterns in why deals are lost (e.g., price, competitor, timing).")

    doc.add_page_break()

    # ── SECTION 9: CONTACTS ───────────────────────────────────
    h1(doc, "9. Contacts Management")
    add_img(doc, "contacts", 6.3, "Contacts List — Individual contact records linked to organisations")
    body(doc,
        "Contacts represent individual people associated with your business relationships. Unlike Leads (which "
        "are unqualified prospects), Contacts are linked to specific Deals or Organizations, providing "
        "a rich, relational view of the people behind each deal.")

    h2(doc, "9.1 Creating a Contact")
    numbered(doc, "Navigate to 'Contacts' in the sidebar.")
    numbered(doc, "Click '+ Add Contact' and fill in First Name, Last Name, Email, and Phone.")
    numbered(doc, "Link the Contact to an Organization and/or a specific Deal.")
    numbered(doc, "Click 'Save'. The contact is now searchable and linkable across the platform.")

    h2(doc, "9.2 Contact Timeline")
    body(doc,
        "Click any contact's name to open their detail view. The Activity Timeline shows a chronological log "
        "of every interaction: calls, notes, and deal stage changes associated with that contact.")

    doc.add_page_break()

    # ── SECTION 10: ORGANIZATIONS ─────────────────────────────
    h1(doc, "10. Organizations & Account Management")
    add_img(doc, "organizations", 6.3, "Organizations List — 15+ corporate accounts from Hyperion Systems to Pinnacle Technologies")
    body(doc,
        "Organizations represent corporate entities — the companies your team works with. By grouping Contacts, "
        "Leads, and Deals under a single Organization, DigiXCrm gives you a 360-degree view of every "
        "institutional relationship.")

    h2(doc, "10.1 Creating an Organization")
    numbered(doc, "Navigate to 'Organizations' in the sidebar.")
    numbered(doc, "Click '+ Create Organization'. Enter the Company Name and Website URL.")
    numbered(doc, "Click 'Save'. The organization appears in your list with its corporate domain displayed.")

    h2(doc, "10.2 Organization Detail View")
    body(doc,
        "Clicking any organization name opens its detail view, showing:")
    bullet(doc, "All Contacts associated with the organization.")
    bullet(doc, "All Leads linked to the organization.")
    bullet(doc, "All Deals (open, won, and lost) associated with the account.")
    tip_box(doc, "Before creating a new lead or deal, always check if the client already exists as an Organization. Linking records to organizations prevents duplicate data and gives you cleaner reporting.")

    doc.add_page_break()

    # ── SECTION 11: PERFORMANCE ───────────────────────────────
    h1(doc, "11. Performance Hub & Quota Tracking")
    add_img(doc, "performance", 6.3, "Performance Hub — $3M+ revenue generated, 250 total leads, 30% conversion rate")
    body(doc,
        "The Performance Hub provides full transparency into how individuals and teams are performing against "
        "their monthly commercial targets. It is visible to all users, creating a culture of healthy, data-driven accountability.")

    h2(doc, "11.1 Overall Team KPIs")
    bullet(doc, "Revenue Generated: Total value of all WON deals in the selected date range.")
    bullet(doc, "Monthly Target: The revenue goal set by the Admin. A dynamic progress bar shows percentage attained.")
    bullet(doc, "Conversion Rate: The ratio of won deals to total leads in the same period.")
    bullet(doc, "Total Leads: The aggregate prospect count visible to the user's scope.")

    h2(doc, "11.2 Revenue Growth Chart")
    body(doc,
        "The area chart visualises cumulative revenue growth over the selected period. Hover over any data point "
        "to see the exact revenue figure for that day.")

    h2(doc, "11.3 Setting Team Targets")
    body(doc,
        "Administrators can set monthly revenue targets directly from the Performance Hub. Targets are "
        "workspace-specific, meaning each business unit can have its own quota benchmarks.")
    numbered(doc, "Scroll to the 'Set Target' panel at the bottom of the Performance Hub.")
    numbered(doc, "Enter the Monthly Revenue Target amount and click 'Save Target'.")
    numbered(doc, "The target bar updates immediately for all users in the workspace.")

    h2(doc, "11.4 Individual Leaderboard")
    body(doc,
        "Below the team KPIs, a ranked leaderboard shows each team member's revenue contribution and lead count. "
        "This is visible to all roles and encourages healthy competition within the sales team.")

    doc.add_page_break()

    # ── SECTION 12: ACTIVITIES ────────────────────────────────
    h1(doc, "12. Activities & Interaction History")
    body(doc,
        "Activities are time-stamped interaction records linked to a Lead, Contact, or Deal. They serve as "
        "the institutional memory of every client relationship in the system.")
    h2(doc, "12.1 Logging an Activity")
    numbered(doc, "Open any Lead, Contact, or Deal detail view.")
    numbered(doc, "Scroll to the 'Activity' or 'Notes' section.")
    numbered(doc, "Select the activity type (Call, Note, Email, Meeting).")
    numbered(doc, "Write your notes in the text area and click 'Save'.")
    body(doc, "All activities are timestamped and attributed to the logged-in user. They appear in reverse-chronological order in the timeline.")

    h2(doc, "12.2 Global Activities Feed")
    body(doc,
        "The 'Activities' item in the sidebar provides a global feed of all activities logged across the "
        "workspace. Filter by user, date range, or activity type to drill into specific interaction patterns.")

    doc.add_page_break()

    # ── SECTION 13: GLOBAL SEARCH ─────────────────────────────
    h1(doc, "13. Global Search")
    body(doc,
        "The Global Search bar (Ctrl+K / Cmd+K) at the top of every page allows you to search across "
        "Leads, Deals, Contacts, and Organizations simultaneously. Results appear instantly as you type, "
        "grouped by entity type for quick navigation.")
    tip_box(doc, "Use Global Search to quickly pull up any record without navigating through menus. It is the fastest way to jump to a specific client or deal mid-conversation.")

    # ── SECTION 14: AUTOMATIONS ───────────────────────────────
    h1(doc, "14. Workflow Automations")
    add_img(doc, "automations", 6.3, "Workflow Engine — Visual automation rule builder in System Settings")
    body(doc,
        "The DigiXCrm Workflow Engine allows Administrators to build 'If This, Then That' automation rules "
        "that eliminate repetitive manual tasks and ensure no lead falls through the cracks.")

    h2(doc, "14.1 How Automations Work")
    body(doc,
        "An automation rule consists of three parts:")
    bullet(doc, "Trigger: The event that starts the automation (e.g., 'Lead Created', 'Deal Won', 'Lead Status Changed').")
    bullet(doc, "Condition (optional): A filter that the trigger must match (e.g., Source = 'PPC').")
    bullet(doc, "Action: What happens automatically (e.g., Assign Owner, Send Email, Send WhatsApp, Create Task).")

    h2(doc, "14.2 Building a Workflow Rule")
    numbered(doc, "Navigate to System Settings > Automations.")
    numbered(doc, "Click '+ Build Workflow'.")
    numbered(doc, "Select a Trigger from the dropdown (e.g., LEAD_CREATED).")
    numbered(doc, "Optionally set a Condition (e.g., Source = 'Cold Call').")
    numbered(doc, "Select an Action (e.g., ASSIGN_OWNER) and set the Action Value (e.g., the Manager's user ID).")
    numbered(doc, "Toggle the rule to 'Active' and click 'Save'.")

    h2(doc, "14.3 Example Automation Scenarios")
    bullet(doc, "VIP PPC Lead Routing: When a lead is created with Source = 'PPC', automatically assign it to the Senior Account Manager.")
    bullet(doc, "New Lead Welcome Email: When any lead is created, trigger a personalised 'Welcome' email from your SMTP integration.")
    bullet(doc, "Deal Won Celebration: When a deal is moved to 'WON', send a WhatsApp message to the team group to celebrate.")
    bullet(doc, "SLA Warning: When a deal is created with a value over $50,000, create an urgent follow-up task within 15 minutes.")

    doc.add_page_break()

    # ── SECTION 15: SETTINGS ──────────────────────────────────
    h1(doc, "15. System Settings & Enterprise Configuration")
    body(doc,
        "System Settings is your master control panel. Access it by clicking 'System Settings' at the "
        "bottom of the sidebar. The settings are tabbed for easy navigation.")

    h2(doc, "15.1 My Profile")
    body(doc,
        "Update your profile details: Full Name, Professional Title, Phone Number, personal Bio, and Timezone. "
        "You can also enable Dark/Light mode, change your preferred language, and configure Two-Factor Authentication (2FA).")

    h2(doc, "15.2 Company & General")
    body(doc,
        "Set your organisation's details: Company Name, Industry, Website, and Timezone. Configure the default "
        "Currency for all deal values in this workspace.")

    h2(doc, "15.3 Team & Permissions (RBAC)")
    body(doc,
        "The Team & Permissions tab is where you manage your people:")
    bullet(doc, "Invite User: Enter an email address and workspace assignment. The user receives an invitation email and sets up their own password.")
    bullet(doc, "Assign Roles: Choose from Admin, Manager, or Rep at the workspace level.")
    bullet(doc, "Custom Roles: Create enterprise PBAC roles with a granular permission matrix. Toggle permissions for: Delete Leads, Export Data, Manage Automations, View AI Reports, etc.")
    bullet(doc, "Suspend & Reactivate: Account Owners can temporarily suspend user access without deleting the account.")

    h2(doc, "15.4 Integrations Panel")
    body(doc,
        "Connect your external tools:")
    bullet(doc, "SMTP Email: Configure your Gmail, Outlook, or custom SMTP for transactional emails.")
    bullet(doc, "Stripe: Link your Stripe account for processing client payments.")
    bullet(doc, "WhatsApp Business API: Configure your WhatsApp Phone ID and Access Token for automated WhatsApp messages.")
    bullet(doc, "Webhook Endpoints: Receive incoming data from WordPress, LinkedIn, Meta Ads, and Upwork.")

    h2(doc, "15.5 Workspace Management")
    body(doc,
        "Account Owners can create additional workspaces within their subscription. Each workspace is a fully "
        "isolated CRM environment with its own team, pipeline stages, and settings.")

    h2(doc, "15.6 Data Export")
    body(doc,
        "Export your leads, deals, contacts, or organizations as a CSV file at any time. The export respects "
        "your current workspace scope and filter settings. Administrators can schedule automatic exports.")

    h2(doc, "15.7 Security & Danger Zone")
    body(doc,
        "The Security tab lets you configure Cloudflare Turnstile CAPTCHA (bot protection for the login page). "
        "The Danger Zone provides irreversible actions such as 'Delete All Leads' or 'Reset Workspace Data' — "
        "these require administrator confirmation and are logged in the Audit Trail.")

    doc.add_page_break()

    # ── SECTION 16: SAAS CONTROL ROOM ────────────────────────
    h1(doc, "16. SaaS Administration — Control Room")
    body(doc,
        "The Control Room is exclusive to the Master Super Admin (platform owner). It is accessible from the "
        "top of the sidebar under 'Platform Control'.")

    h2(doc, "16.1 Platform Accounts")
    body(doc,
        "View a global directory of all organisations registered on the platform. Each entry shows the "
        "subscription tier, seat count, payment provider, and account health status.")

    h2(doc, "16.2 Organisation Actions")
    bullet(doc, "View all workspaces belonging to an organisation.")
    bullet(doc, "Suspend or reactivate an entire organisation (blocks all users from that account).")
    bullet(doc, "Downgrade or upgrade an organisation's plan tier manually.")
    bullet(doc, "Force a billing refresh or manually mark an account's payment as verified.")

    h2(doc, "16.3 Global Audit Logs")
    body(doc,
        "Every significant action taken on the platform — lead deletion, user invitations, workspace creation, "
        "integration changes — is written to an immutable Audit Log. Super Admins can search and filter "
        "this log by user, action type, entity, date range, and workspace.")

    doc.add_page_break()

    # ── SECTION 17: TROUBLESHOOTING ───────────────────────────
    h1(doc, "17. Troubleshooting & FAQ")

    h2(doc, "Q: I can't see leads created by my colleague.")
    body(doc, "A: Confirm you are in the correct Workspace (check the sidebar switcher). Also verify your role — Reps can only see leads assigned to them, while Managers and Admins can see all workspace leads.")

    h2(doc, "Q: My password reset link says 'already used'.")
    body(doc, "A: Reset links are one-time-use only and expire after 1 hour. Return to the login page and request a new reset link from the 'Forgot your password?' option.")

    h2(doc, "Q: Real-time updates are not appearing without a page refresh.")
    body(doc, "A: DigiXCrm uses Server-Sent Events (SSE) for live updates. Ensure your browser tab is not inactive or throttled. Check if your network or proxy blocks persistent HTTP connections. If the issue persists, the polling fallback will sync data every 30 seconds.")

    h2(doc, "Q: My CSV import shows 0 leads imported.")
    body(doc, "A: Ensure your CSV file has 'firstName' and 'lastName' as column headers (case-sensitive). Remove any special characters from the header row. The file must be saved in UTF-8 encoding.")

    h2(doc, "Q: How do I change the default pipeline stages?")
    body(doc, "A: Go to System Settings > Workspace Settings > Deal Stages. You can add, rename, or reorder stages. Note: changing stages does not affect existing deals already in those stages.")

    h2(doc, "Q: An automation rule is not triggering.")
    body(doc, "A: Verify the rule is toggled to 'Active' in the Automations tab. Check that the Trigger, Condition, and Action values are correctly configured. Review the Audit Log for any automation errors.")

    doc.add_page_break()

    # ── SECTION 18: GLOSSARY ──────────────────────────────────
    h1(doc, "18. Glossary of Terms")
    terms = [
        ("Lead",         "An unqualified prospect — a person or company showing potential interest in your product or service."),
        ("Deal",         "A qualified opportunity actively progressing through your sales pipeline. Created by 'promoting' a Lead."),
        ("Contact",      "An individual person record, typically linked to an Organization, Lead, or Deal."),
        ("Organization", "A corporate entity (company) that groups related Contacts, Leads, and Deals."),
        ("Workspace",    "An isolated CRM environment — like a separate business unit — with its own team, data, and settings."),
        ("Pipeline",     "The visual representation of all open Deals arranged by their stage in the sales lifecycle."),
        ("Stage",        "A phase in the sales lifecycle (e.g., Contacted, Negotiation, WON). Customisable per workspace."),
        ("RBAC",         "Role-Based Access Control. The system that determines what each user can see and do, based on their assigned role."),
        ("PBAC",         "Permission-Based Access Control. A more granular system using custom roles with individually toggled permissions."),
        ("SSE",          "Server-Sent Events. The technology that delivers real-time push updates from the server to connected browsers."),
        ("Audit Log",    "An immutable record of every significant action taken by users in the system, used for security and compliance."),
        ("Automation",   "A configured rule that triggers a predefined action automatically when a specific event occurs in the CRM."),
        ("KPI",          "Key Performance Indicator. A measurable value that demonstrates how effectively a team is achieving key objectives."),
        ("Conversion Rate", "The percentage of leads that have been successfully promoted to deals. A key measure of pipeline quality."),
        ("WON",          "A deal stage indicating the sale has been successfully closed and revenue has been captured."),
        ("ONBOARDED",    "An order status indicating the client has been fully delivered to and set up post-sale."),
    ]
    for term, definition in terms:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{term}: ")
        r1.bold = True
        r1.font.color.rgb = PRIMARY
        r2 = p.add_run(definition)
        r2.font.color.rgb = RGBColor(50, 50, 50)
        p.paragraph_format.space_after = Pt(4)

    # ── FOOTER PAGE ───────────────────────────────────────────
    doc.add_page_break()
    for _ in range(8): doc.add_paragraph("")
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ep.add_run("DigiXCrm CRM  ·  Enterprise User & Administration Guide")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = PRIMARY

    ep2 = doc.add_paragraph()
    ep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ep2.add_run("© 2026 DigiCare Products. All Rights Reserved. Confidential Business Property.")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    ep3 = doc.add_paragraph()
    ep3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ep3.add_run("Support: digicarehouse.sales@gmail.com  ·  Version: 2.4 (Stable)")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

    out = "DigiSales_Professional_User_Manual_v3.docx"
    doc.save(out)
    return os.path.abspath(out)

if __name__ == "__main__":
    path = create_manual()
    print(f"DONE: {path}")
